from typing import List
import os
from docx import Document as DocxDocument

from loaders.base_loader import BaseLoader
from loaders.document import Document


class DocxLoader(BaseLoader):

    def load(self, file_path: str) -> List[Document]:

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.lower().endswith(".docx"):
            raise ValueError("DocxLoader only supports .docx files")
        
        try:
            doc = DocxDocument(file_path)
            
        except Exception as e:
            raise RuntimeError(f"Error reading file: {e}")
        
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
            
        content = "\n".join(paragraphs)


        metadata = {
            "source": file_path,
            "file_type": "docx"
        }
        document = Document(content=content, metadata=metadata)
        return [document]